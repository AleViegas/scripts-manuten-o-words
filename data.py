import os
import time
from colorama import init, Fore
from docx import Document

#Função que substitui a ART do memorial
def replaceDateOnMemorial(path, dia, mes):
    doc = Document(fr"{path}")
    #abrindo o documento
    
    print(doc.paragraphs[0].text)
    texto = f"                                                                                                                 São Paulo, {dia} de {mes} de 2021"
    
    doc.paragraphs[0].text = texto

    doc.save(fr"{path}")
    #Salvando o documento no mesmo caminho
    
    
#Função que substitui a ART da carta
def replaceDateOnLetter(path, dia, mes):
    doc = Document(fr"{path}")
    #abrindo o documento

    texto = f"São Jose dos Campos, {dia} de {mes} de 2021"
    #print(doc.paragraphs[3].text)

    doc.paragraphs[3].text = texto

    #print(doc.paragraphs[3].style)

    doc.save(fr"{path}")
    #Salvando o documento no mesmo caminho

init()
#Colorama init

print(Fore.MAGENTA + "Bem vindo!")
time.sleep(0.5)
print("Você acabou de entrar no Substituidor de ARTs Miralt")
time.sleep(1.5)
dir = input("Cole aqui o diretório principal: ")
time.sleep(1)
dia = input("Passa o dia meu chapa: ")
time.sleep(1)
mes = input("Agora o mes: ")
time.sleep(1)
print("Substituindo...\n")

unsortedSubfolders = [f.path for f in os.scandir(dir) if f.is_dir()]
subfolders = sorted(unsortedSubfolders, key=len)
# f = folder
#Pegando todos os sub diretorios do input


for subfolder in subfolders:
    #Loop for in para substituir a art em todas as pastas (Trechos)
    
    docxs = [f.path for f in os.scandir(subfolder) if f.name.endswith("docx")]
    #Separação dos arquivos .docx
    
    slash = docxs[0].count("\\")
    #Quantidade de \s
    #Vai ser utilizada para o spit
    
    for docx in docxs:
        
        docxlist = docx.split("\\")
        docxName = docxlist[slash]
        #Obtenção do nome do arquivo
        
        if docxName.startswith("Carta"):
            #Separação das Cartas
            
            replaceDateOnLetter(docx, dia, mes)
            #Chamada da função
            
            #print(docxName)
            
        if docxName.startswith("Memorial"):
            #Separação dos memoriais
            
            replaceDateOnMemorial(docx, dia, mes)
            #Chamada da função
            
            #print(docxName)
            
        
print("\nA ART foi adicionada em todos os documentos")
time.sleep(1)
print("Bora pra cima!")
time.sleep(5)