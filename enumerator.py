import os
import time
from colorama import init, Fore
from docx import Document
from docx.shared import Pt
import pandas as pd
  
#Função que numera as cartas
def enumerateLetter(path):
    doc = Document(fr"{path}")

    table = doc.tables[0]
    #Selecionando a segunda tabela do documento

    inftrecho = np[i]

    carta = f"Número do Projeto: {inftrecho[1]}/2021"

    doc.paragraphs[0].text = []
    run = doc.paragraphs[0].add_run(carta)
    font = run.font
    font.name = 'Calibri (Body)'
    font.size = Pt( 9 )

    doc.save(fr"{path}")
    #Salvando o documento no mesmo caminho

init()
#Colorama init

print(Fore.MAGENTA + "ENUMERATOR!")
#time.sleep(0.5)
print("Você acabou de entrar no ENUMERATOR de ARTs Miralt")
#time.sleep(1.5)
dir = input("Cole aqui o diretório principal: ")
#time.sleep(1)
excel = input("Passa o logradouros meu chapa (só o nome): ")
#time.sleep(1)
print("Substituindo...\n")

unsortedSubfolders = [f.path for f in os.scandir(dir) if f.is_dir()]
subfolders = sorted(unsortedSubfolders, key=len)
#print(subfolders)
# f = folder
#Pegando todos os sub diretorios do input

excel = f"{excel}.xlsx"
df = pd.read_excel(excel, index_col=0)
#instancia do excel

np = df.to_numpy()
#trasnformando o excel em np array

i = 0
#control

for subfolder in subfolders:
    #Loop for in para substituir a art em todas as pastas (Trechos)
    
    docxs = [f.path for f in os.scandir(subfolder) if f.name.endswith("docx")]
    #Separação dos arquivos .docx
    
    slash = docxs[0].count("\\")
    #Quantidade de \s
    #Vai ser utilizada para o spit

    i = i + 1
    
    for docx in docxs:
        
        docxlist = docx.split("\\")
        docxName = docxlist[slash]
        #Obtenção do nome do arquivo
        
        if docxName.startswith("Carta"):
            #Separação das Cartas
            
            enumerateLetter(docx)
            #Chamada da função
            
            print(docxName)
            

        
print("\nEMUMERATORRRR")
time.sleep(1)
print("Bora pra cima!")
time.sleep(5)