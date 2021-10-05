import os
import time
from colorama import init, Fore
from docx import Document
from docx.shared import Pt

#Função que substitui a ART do memorial
def replaceArtOnMemorial(path, artProj, artExec):
    doc = Document(fr"{path}")
    #abrindo o documento
    
    fontName = "Times New Roman"
    fontSize = 12

    text = f"{artProj}\n{artExec}"
    #Texto que sera colocado no documento

    table = doc.tables[0]
    #Selecionando a primeira tabela do documento

    table.cell(0, 17).text = []
    #Limpando a Celula

    para = table.cell(0, 17).paragraphs[0]
    run = para.add_run(text)
    #Adicionando a run com texto

    font = run.font
    font.name = fontName
    font.size = Pt( fontSize )
    #Style da run    

    doc.save(fr"{path}")
    #Salvando o documento no mesmo caminho
    
    
def replaceArtOnLetter(path, art, celposition):
    doc = Document(fr"{path}")
    #abrindo o documento

    fontName = "Calibri (Body)"
    fontSize = 9
    
    table = doc.tables[0]
    #Selecionando a segunda tabela do documento
    
    textList = table.cell(celposition[0], celposition[1]).text.split("\n")
    #Separando todo o texto da celula em uma lista
    
    textList[0] = f"ART - {art}"
    #Deletando o ultimo texto (ART)
    
    barraN = "\n"
    text = barraN.join(textList)
    #Juntando toda a lista
    #Colocando um espaço entre cada termo

    table.cell(celposition[0], celposition[1]).text = []
    #Limpando a Celula

    para = table.cell(celposition[0], celposition[1]).paragraphs[0]
    run = para.add_run(text)
    #Adicionando a run com texto

    font = run.font
    font.name = fontName
    font.size = Pt( fontSize )
    #Style da run    
    
    doc.save(fr"{path}")
    #Salvando o documento no mesmo caminho

init()
#Colorama init

print(Fore.MAGENTA + "Bem vindo!")
time.sleep(0.5)
print("Você acabou de entrar no movimento artistico da Miralt")
time.sleep(1.5)
dir = input("Cole aqui o diretório principal: ")
time.sleep(1)
artProj = input("Passa a Art de Projeto meu chapa: ")
time.sleep(1)
artExec = input("agora a Art de execução: ")
time.sleep(1)
print("Substituindo...\n")

unsortedSubfolders = [f.path for f in os.scandir(dir) if f.is_dir()]
subfolders = sorted(unsortedSubfolders, key=len)
# f = folder
#Pegando todos os sub diretorios do input


for subfolder in subfolders:
    #Loop for in para substituir a art em todas as pastas (Trechos)
    
    docxs = [f.path for f in os.scandir(subfolder) if f.name.endswith("docx")]
    #Separação dos arquivos .docx
    
    slash = docxs[0].count("\\")
    #Quantidade de \s
    #Vai ser utilizada para o spit
    
    for docx in docxs:
        
        docxlist = docx.split("\\")
        docxName = docxlist[slash]
        #Obtenção do nome do arquivo
        
        if docxName.startswith("Carta"):
            #Separação das Cartas
            
            replaceArtOnLetter(docx, artProj, [1, 4])
            replaceArtOnLetter(docx, artExec, [3, 4])

            #Chamada da função
            
            print(docxName)
            
        if docxName.startswith("Memorial"):
            #Separação dos memoriais
            
            replaceArtOnMemorial(docx, artProj, artExec)
            #Chamada da função
            
            print(docxName)
            
        
print("\nA ART foi adicionada em todos os documentos")
time.sleep(1)
print("Bora pra cima!")
time.sleep(5)